"""Generación de documentos Word con el feedback estructurado de cada entrega."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from core.evaluator import Evaluacion

_COLOR_HEADER = RGBColor(0x1F, 0x3A, 0x5F)
_COLOR_MUTED = RGBColor(0x5A, 0x5A, 0x5A)

_COMPLETITUD_COLORS = {
    "Completo": RGBColor(0x1E, 0x7D, 0x32),
    "Parcial": RGBColor(0xB8, 0x86, 0x0B),
    "Insuficiente": RGBColor(0xC6, 0x3A, 0x2B),
    "No realizado": RGBColor(0x8A, 0x8A, 0x8A),
}


def _shade_cell(cell, color_hex: str) -> None:
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {qn("w:fill"): color_hex})
    shading.append(shd)


def generar_docx_feedback(evaluacion: Evaluacion, original_filename: str) -> BytesIO:
    """Genera un .docx en memoria con el feedback estructurado de una entrega."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    title = doc.add_heading("Retroalimentación de Evaluación", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = _COLOR_HEADER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Entrega original: {original_filename}\n"
        f"Fecha de corrección: {datetime.now().strftime('%d-%m-%Y %H:%M')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = _COLOR_MUTED

    doc.add_paragraph()

    doc.add_heading("Resumen General", level=1)
    doc.add_paragraph(evaluacion.resumen_general)

    doc.add_heading("Evaluación por Criterio", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    headers = ["Criterio", "Puntaje", "Completitud", "Feedback"]
    for cell, text in zip(header_cells, headers):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        _shade_cell(cell, "1F3A5F")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for criterio in evaluacion.criterios:
        row_cells = table.add_row().cells

        row_cells[0].text = criterio.nombre_criterio

        row_cells[1].text = f"{criterio.puntaje_obtenido:g} / {criterio.puntaje_maximo:g}"

        completitud_run = row_cells[2].paragraphs[0].add_run(criterio.grado_completitud)
        completitud_run.bold = True
        completitud_run.font.color.rgb = _COMPLETITUD_COLORS.get(
            criterio.grado_completitud, _COLOR_MUTED
        )

        row_cells[3].text = criterio.feedback

    doc.add_paragraph()

    doc.add_heading("Resultado Final", level=1)
    resultado = doc.add_table(rows=2, cols=2)
    resultado.style = "Light List Accent 1"
    resultado.rows[0].cells[0].text = "Puntaje total"
    resultado.rows[0].cells[1].text = (
        f"{evaluacion.puntaje_total:g} / {evaluacion.puntaje_maximo_total:g}"
    )
    resultado.rows[1].cells[0].text = "Nota sugerida"
    resultado.rows[1].cells[1].text = f"{evaluacion.nota_sugerida:.1f}"

    for row in resultado.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else None

    doc.add_paragraph()
    nota_final = doc.add_paragraph()
    nota_final.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_run = nota_final.add_run(
        "Este documento es editable: el profesor puede ajustar puntajes, comentarios "
        "o la nota sugerida antes de entregarlo al estudiante."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = _COLOR_MUTED

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def feedback_filename(original_filename: str) -> str:
    stem = original_filename.rsplit(".", 1)[0]
    return f"Feedback_{stem}.docx"
