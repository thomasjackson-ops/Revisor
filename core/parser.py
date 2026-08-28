"""Extracción de texto desde archivos .pdf, .docx, .csv y .xlsx de forma robusta."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx"}
SUPPORTED_SPREADSHEET_EXTENSIONS = {".csv", ".xlsx", ".xls"}
SUPPORTED_EXTENSIONS = SUPPORTED_DOCUMENT_EXTENSIONS | SUPPORTED_SPREADSHEET_EXTENSIONS


class UnsupportedFileTypeError(Exception):
    """Se lanza cuando la extensión del archivo no está soportada."""


class FileExtractionError(Exception):
    """Se lanza cuando el archivo no pudo leerse (corrupto, protegido, vacío, etc.)."""


@dataclass
class ExtractedDocument:
    """Resultado de la extracción de texto de una entrega."""

    filename: str
    text: str
    char_count: int


def extract_text(file_path: Path) -> ExtractedDocument:
    """Extrae el texto/contenido de un archivo .pdf, .docx, .csv o .xlsx.

    Lanza UnsupportedFileTypeError o FileExtractionError en caso de fallo,
    para que el llamador pueda saltar el archivo sin detener el lote completo.
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        text = _extract_pdf(file_path)
    elif suffix == ".docx":
        text = _extract_docx(file_path)
    elif suffix == ".csv":
        text = _extract_csv(file_path)
    elif suffix in (".xlsx", ".xls"):
        text = _extract_excel(file_path)
    else:
        raise UnsupportedFileTypeError(f"Tipo de archivo no soportado: {suffix}")

    text = text.strip()
    if not text:
        raise FileExtractionError(
            "No se pudo extraer texto/datos legibles del archivo (puede ser una imagen escaneada, "
            "una planilla vacía o un formato dañado)."
        )

    return ExtractedDocument(filename=file_path.name, text=text, char_count=len(text))


def _extract_pdf(file_path: Path) -> str:
    """Intenta extraer texto con pdfplumber; si falla, recurre a pypdf como respaldo."""
    try:
        chunks: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    chunks.append(page_text)
        text = "\n".join(chunks)
        if text.strip():
            return text
    except Exception:
        pass

    try:
        reader = PdfReader(str(file_path))
        chunks = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(chunks)
    except Exception as exc:
        raise FileExtractionError(f"Error al leer PDF con pypdf: {exc}") from exc


def _extract_docx(file_path: Path) -> str:
    try:
        doc = DocxDocument(str(file_path))
    except Exception as exc:
        raise FileExtractionError(f"Error al abrir el .docx: {exc}") from exc

    parts: list[str] = [p.text for p in doc.paragraphs if p.text]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts)


def _dataframe_to_text(df: pd.DataFrame, max_rows: int = 500) -> str:
    """Convierte un DataFrame a texto tabular legible por el LLM."""
    if df.empty:
        return "(planilla sin filas de datos)"

    truncated = df.shape[0] > max_rows
    if truncated:
        df = df.head(max_rows)

    text = df.to_string(index=False, na_rep="")
    if truncated:
        text += f"\n\n[... se truncó la vista a las primeras {max_rows} filas de {df.shape[0]} totales ...]"
    return text


def _extract_csv(file_path: Path) -> str:
    # dtype=str evita que pandas "adivine" tipos numéricos y corrompa datos de
    # contacto (teléfonos con "+", ceros a la izquierda, IDs largos, etc.).
    try:
        df = pd.read_csv(file_path, dtype=str, keep_default_na=False)
    except Exception:
        try:
            df = pd.read_csv(file_path, sep=";", dtype=str, keep_default_na=False)
        except Exception as exc:
            raise FileExtractionError(f"Error al leer CSV: {exc}") from exc

    return _dataframe_to_text(df)


def _extract_excel(file_path: Path) -> str:
    try:
        sheets = pd.read_excel(file_path, sheet_name=None, dtype=str, keep_default_na=False)
    except Exception as exc:
        raise FileExtractionError(f"Error al leer Excel: {exc}") from exc

    parts: list[str] = []
    for sheet_name, df in sheets.items():
        parts.append(f"--- Hoja: {sheet_name} ---")
        parts.append(_dataframe_to_text(df))

    return "\n\n".join(parts)
